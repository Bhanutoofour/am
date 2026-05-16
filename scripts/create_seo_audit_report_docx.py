from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = "reports/Autocracy_SEO_Audit_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE5", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_border(section, color="FFC400"):
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "18")
        border.set(qn("w:color"), color)
        pg_borders.append(border)
    sect_pr.append(pg_borders)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(18 if level == 1 else 14)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def style_header_cells(cells):
    for cell in cells:
        set_cell_shading(cell, "0B0F14")
        set_cell_margins(cell, 120, 140, 120, 140)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.2)
    set_table_borders(table)
    headers = table.rows[0].cells
    headers[0].text = "Audit Area"
    headers[1].text = "Current Status"
    style_header_cells(headers)

    rows = [
        ("Primary domain", "https://www.autocracymachinery.com remains the preferred site URL."),
        ("Non-www redirect", "https://autocracymachinery.com redirects permanently to the www domain."),
        ("Canonicals", "Listing page canonicals have been added for /products and /industries."),
        ("Legacy URLs", "Old query URLs now redirect to clean product or industry routes."),
        ("Sitemap freshness", "Products, industries, and models now expose updatedAt values and sitemap uses those dates."),
        ("Performance baseline", "Remote images can be served as AVIF/WebP, key third-party origins have preconnect hints, and Meta Pixel loads on idle."),
    ]

    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, 140, 160, 140, 160)
    doc.add_paragraph()


def add_status_table(doc):
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [0.9, 1.1, 2.1, 2.3, 1.15]
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
    set_table_borders(table)

    headers = ["Priority", "Area", "Original Issue", "Resolution Applied", "Status"]
    for index, heading in enumerate(headers):
        table.rows[0].cells[index].text = heading
    style_header_cells(table.rows[0].cells)

    rows = [
        (
            "High",
            "Homepage metadata",
            "Homepage title had typo: Autocracy Mchinery.",
            "Corrected title, Open Graph title, and Twitter title to Autocracy Machinery.",
            "Resolved",
        ),
        (
            "High",
            "Homepage H1",
            "Global homepage used India-focused hidden H1 text.",
            "Global / now uses global H1 wording; /en-in keeps India-focused wording.",
            "Resolved",
        ),
        (
            "High",
            "Canonical tags",
            "/products was missing a canonical tag.",
            "Added canonical plus en-IN and x-default alternates.",
            "Resolved",
        ),
        (
            "High",
            "Canonical tags",
            "/industries was missing a canonical tag.",
            "Added canonical plus en-IN and x-default alternates.",
            "Resolved",
        ),
        (
            "High",
            "Duplicate URLs",
            "Old query URLs returned 200.",
            "Added permanent redirects for productId and industryId query URLs on global and en-in listing pages.",
            "Resolved",
        ),
        (
            "Medium",
            "Legacy model URLs",
            "Legacy /product/... URLs should redirect to clean model URLs.",
            "Confirmed legacy route uses permanentRedirect to /products/{category}/{model}.",
            "Resolved",
        ),
        (
            "Medium",
            "Sitemap freshness",
            "Product, industry, and model sitemap dates are static.",
            "Added createdAt/updatedAt fields, applied the Neon migration, and wired app/sitemap.ts to use updatedAt values.",
            "Resolved",
        ),
        (
            "Medium",
            "Blog H1",
            "Blog page H1 was generic: Blogs.",
            "Changed H1 to Machinery Blogs and Industry Insights.",
            "Resolved",
        ),
        (
            "Medium",
            "Industry keywords",
            "Industry pages had short hidden H1 wording.",
            "Improved hidden SEO H1 with machinery/equipment solution phrasing.",
            "Resolved",
        ),
        (
            "Low",
            "Performance SEO",
            "Third-party scripts and large images may affect mobile speed.",
            "Enabled AVIF/WebP image optimization for configured remote hosts, added preconnect hints, and moved Meta Pixel to lazyOnload.",
            "Resolved",
        ),
    ]

    status_colors = {
        "Resolved": "E7F4E8",
    }

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[index], 120, 120, 120, 120)
            if index == 4:
                set_cell_shading(cells[index], status_colors.get(value, "FFFFFF"))
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()


def create_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    add_page_border(section)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Arial"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Autocracy Machinery Website SEO Audit")
    title_run.bold = True
    title_run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Resolved SEO status after code and database updates")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("Updated: May 15, 2026 | Domain: www.autocracymachinery.com")
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()
    add_heading(doc, "Executive Summary", 1)
    p = doc.add_paragraph()
    p.add_run(
        "The high-priority SEO issues from the audit have been resolved in code. "
        "The site now has corrected homepage metadata, improved global/India page signals, "
        "canonical tags on the product and industry listing pages, permanent redirects for old query URLs, "
        "dynamic sitemap freshness through CMS updatedAt fields, and a small performance baseline pass for image and third-party script loading."
    )

    add_summary_table(doc)

    add_heading(doc, "Resolved Fix Status", 1)
    add_status_table(doc)

    add_heading(doc, "Code Areas Updated", 1)
    for item in [
        "Homepage metadata, hidden H1 signals, blog H1, and industry SEO headings were corrected.",
        "Product and industry listing pages now include canonicals, hreflang alternates, and legacy query redirects.",
        "Product, industry, and model records now support createdAt/updatedAt timestamps for sitemap freshness.",
        "public/llms.txt now uses the preferred www domain, and the performance baseline includes image optimization, preconnect hints, and idle Meta Pixel loading.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Remaining Recommendation", 1)
    p = doc.add_paragraph()
    p.add_run(
        "After deployment, submit the refreshed sitemap in Google Search Console and monitor PageSpeed Insights. "
        "If PageSpeed flags any specific page images, replace those assets with right-sized designer exports or add responsive CMS variants."
    )

    add_heading(doc, "Verification", 1)
    p = doc.add_paragraph()
    p.add_run(
        "npm run build passed. The Neon timestamp migration was applied. SITE_URL drives the www sitemap and robots configuration, and legacy /product/[slug] redirects to clean model URLs."
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Autocracy Machinery SEO Audit - Resolved Status | May 2026")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)

    import os

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    create_document()
