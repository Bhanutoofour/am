from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("outputs/image-resolution-inventory")
OUTPUT_DOCX = OUTPUT_DIR / "autocracy-image-resolution-inventory.docx"

YELLOW = "F7C300"
BLACK = "01060A"
LIGHT = "F5F5F5"
MID = "D7D7D7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=MID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:{}".format(margin)
        element = margins.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_paragraph(paragraph, size=9, bold=False, color=BLACK, before=0, after=0, line_spacing=1.05):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLACK)
    run.font.size = Pt(18 if level == 1 else 13)
    return paragraph


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF6D4")
    set_cell_border(cell, YELLOW, "10")
    set_cell_margins(cell, 140, 180, 140, 180)
    paragraph = cell.paragraphs[0]
    paragraph.add_run(text)
    style_paragraph(paragraph, size=9.5, bold=False, color=BLACK, line_spacing=1.15)


def add_table(doc, title, headers, rows, widths, page_break_before=False):
    if page_break_before:
        doc.add_page_break()
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, widths)

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for cell, header in zip(header_row.cells, headers):
        set_cell_shading(cell, BLACK)
        set_cell_border(cell, BLACK)
        set_cell_margins(cell, 120, 110, 120, 110)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run(header)
        style_paragraph(paragraph, size=8.25, bold=True, color="FFFFFF", line_spacing=1.05)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for index, row_data in enumerate(rows):
        cells = table.add_row().cells
        for cell_index, value in enumerate(row_data):
            cell = cells[cell_index]
            if index % 2 == 0:
                set_cell_shading(cell, "FFFFFF")
            else:
                set_cell_shading(cell, LIGHT)
            set_cell_border(cell)
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.add_run(value)
            style_paragraph(paragraph, size=8, color=BLACK, line_spacing=1.05)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    return table


def build_doc():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Autocracy Machinery")
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(YELLOW)

    headline = doc.add_paragraph()
    headline.paragraph_format.space_after = Pt(8)
    run = headline.add_run("Website Image Resolution Inventory")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLACK)

    intro = doc.add_paragraph()
    intro.add_run(
        "Designer handoff for replacing website images. Sizes are based on current Next.js components, CMS upload fields, and CSS containers."
    )
    style_paragraph(intro, size=10.5, color="333333", after=8, line_spacing=1.2)

    add_note(
        doc,
        "General rule: use JPG/WebP for real photographs and PNG/WebP with transparent background for isolated machine cutouts. "
        "For cover areas, keep the subject in the center-safe zone. For contain areas, keep the full machine visible with clean breathing space."
    )

    cms_rows = [
        ("Home hero slider image", "Homepage top carousel", "Hero CMS: image; folder hero", "Full-width banner, 581px desktop height, object-fit cover", "Min 1920 x 581; preferred 2400 x 726; aspect 3.3:1"),
        ("Industry thumbnail", "Homepage Choose Your Industry, /industries, related cards", "industries.thumbnail; folder industries/thumbnails", "Card image, desktop height 320px, mobile 200px, object-fit cover", "Preferred 1200 x 1000; minimum 800 x 800"),
        ("Industry page banner", "/industries/[industry] hero slider", "industries.bannerImages[].imageUrl; folder industries/banners", "Wide banner, max 1280 x 440, object-fit cover", "Min 1280 x 440; preferred 2560 x 880; aspect 2.91:1"),
        ("Product category thumbnail", "Product cards on homepage, /products, industry product lists, mega menus", "products.thumbnail; folder products/thumbnails", "Mostly object-fit contain; desktop card area about 245 x 260", "Preferred 1120 x 720; minimum 560 x 360"),
        ("Product general image", "Product detail intro fallback and some model sections", "products.generalImage; folder products/general", "Contained product image, no crop", "Preferred 1120 x 720 or 1500 x 900"),
        ("Product SEO/social image", "Google/social preview image", "products.seoSocialImage; folder products/social", "Metadata only", "1200 x 630"),
        ("Model thumbnail", "Product model cards, model carousels, mobile model cards", "models.thumbnail; folder models/thumbnails", "Desktop visual area about 300 x 118; mobile source 500 x 250; object-fit contain", "Preferred 1000 x 500; minimum 500 x 250"),
        ("Model cover image", "/product/[model] and industry model detail hero/media slider", "models.coverImage; folder models/covers", "Main model media, often 1500 x 768 wide hero or 16:10, object-fit cover", "Preferred 3000 x 1536; minimum 1500 x 768"),
        ("Model detail / gallery image", "Model media carousel and industry application blocks", "models.modelDescription[].image; folder models/details", "Extra carousel/detail images; 16:9 or 16:10 depending section", "Hero/gallery 3000 x 1536; application 1536 x 864"),
        ("Model brochure", "Brochure buttons/downloads", "models.brochure; folder models/brochures", "File download, not a photo slot", "PDF"),
        ("Industry brochure", "Industry brochure downloads", "industries.brochure; folder industries/brochures", "File download, not a photo slot", "PDF"),
        ("Blog banner", "/blogs, /blog/[slug], social preview", "blogs.banner; folder blogs/banners", "Blog cards use 16:9; detail max height about 342px; object-fit cover", "Preferred 1600 x 900; minimum 1200 x 675"),
        ("Blog SEO/social image", "Blog OpenGraph/Twitter preview", "blogs.seoSocialImage", "Metadata only", "1200 x 630"),
    ]

    static_rows = [
        ("Home FAQ CTA image", "Homepage FAQ + lead form section", "INDUSTRY.SAMPLE_INDUSTRY -> assets/hero_section/multi-chain-trencher.png", "Left container min-height 448px desktop, 352px mobile, object-fit contain", "Preferred 1520 x 1040; minimum 760 x 520"),
        ("Recognition awards carousel photos", "Homepage/About recognition section", "IMAGES.RECOGNITION.AWARDS.AWARD_07/08", "Large carousel image 704 x 344 desktop, 353 x 178 mobile, object-fit cover", "Preferred 1408 x 688; minimum 704 x 344; aspect 2.05:1"),
        ("Award logos", "Recognition award cards", "data/recognitionsData.ts -> assets/recognitions/Awards/*.svg", "Logo card image around 118 x 118, object-fit contain", "SVG preferred; PNG fallback 300 x 300"),
        ("Certificate logos", "Our Certifications section", "data/recognitionsData.ts -> assets/recognitions/certificates/*.svg", "Rendered around 150 x 150, object-fit contain", "SVG preferred; PNG fallback 300 x 300"),
        ("Media/news logos", "Recognition media cards", "data/recognitionsData.ts -> assets/recognitions/media/*.svg", "Small contained logo/card image", "SVG preferred; PNG fallback 300 x 300"),
        ("Client logos", "Testimonials/client logo strip", "constants/Images/images.ts -> assets/recognitions/clients/*.svg", "Contained logo strip, max height about 70px", "SVG preferred; PNG fallback 480 x 140"),
        ("Header logo", "Header/nav", "IMAGES.LOGO -> assets/icons/logo.svg", "Around 162px wide by 37px high", "SVG preferred; current size is fine"),
        ("Footer logo", "Footer", "IMAGES.LOGO", "Around 170px by 40px", "SVG preferred; current size is fine"),
    ]

    page_rows = [
        ("/ homepage hero", "CMS hero image", "component/sections/caraousel/Caraousel.tsx; component/sections/caraousel/styles.module.scss", "Banner crops to fill. Avoid important details at far edges."),
        ("/ Choose Your Industry", "CMS industry thumbnail", "component/sections/Industries/Industries.tsx; component/molecules/industryCard/IndustryCard.tsx", "Cards crop on desktop/mobile. Put subject in center 70%."),
        ("/ Our Products Lineups", "CMS product thumbnail", "component/sections/products/Products.tsx; component/molecules/productCard/ProductCard.tsx", "Use isolated product cutouts with white/transparent background."),
        ("/ FAQ CTA", "Static placeholder machine image", "app/(main)/HomeFaqCta.tsx", "Replace SAMPLE_INDUSTRY if a stronger CTA machine photo is needed."),
        ("/products listing", "CMS product thumbnail", "app/(main)/products/ProductsClient.tsx; component/molecules/productCard/ProductCard.tsx", "Same product card rules. Leave enough white space around all machines."),
        ("/products/[slug] product top image", "CMS product thumbnail or generalImage", "app/(main)/products/[slug]/ProductClient.tsx", "Displayed at 560 x 360 source, but provide 1120 x 720 for retina/zoom."),
        ("/product/[model] model main media", "CMS model coverImage plus detail images", "app/(main)/product/[slug]/ProductModalClient.tsx; modalStyles.module.scss", "Most important model photo. Wide field/action photo works best."),
        ("/industries/[slug] industry hero", "CMS industry bannerImages", "app/(main)/industries/[slug]/IndustryClient.tsx", "Use clean wide banner with product/site context visible after cropping."),
        ("/industries/[slug]/[productSlug]", "CMS product/model images", "industry product routes", "Same product/model image rules apply."),
        ("/blog and /blog/[slug]", "CMS blog banner", "app/(main)/blog/BlogsClient.tsx; app/(main)/blog/[slug]", "Export 16:9 feature images; avoid text-heavy graphics."),
        ("/brochure cards", "Product/model brochure cover or thumbnails", "component/molecules/brochureCard/BrochureCard.tsx", "If creating brochure covers, use 200 x 282 A-series ratio."),
        ("/hire-rental-industry-equipment", "Rental equipment thumbnail", "app/(main)/hire-rental-industry-equipment/HireEquipmentClient.tsx", "Rendered about 496 x 335 desktop; export 992 x 670."),
        ("/about-us imagery", "Static/about page images", "app/(main)/about-us/AboutUsClient.tsx; aboutStyles.module.scss", "Mixed cover images. Use high-quality 16:9 or square images depending block."),
        ("/careers imagery", "Static careers images", "app/(main)/careers/CareersClient.tsx; careersStyles.module.scss", "Mixed decorative photos. Export at least 2x declared component sizes."),
    ]

    checklist_rows = [
        ("1", "Hero/banner images", "Homepage 2400 x 726; industry banners 2560 x 880"),
        ("2", "Product/model machine cutouts", "Product cards 1120 x 720; model cards 1000 x 500"),
        ("3", "Model detail hero/gallery photos", "3000 x 1536"),
        ("4", "Industry thumbnails", "1200 x 1000, subject centered"),
        ("5", "Blog images", "1600 x 900 and social image 1200 x 630"),
        ("6", "Recognition/certificate/client logos", "SVG preferred; PNG fallback at 2x display size"),
    ]

    add_table(
        doc,
        "CMS Upload Slots",
        ["Image slot", "Where it appears", "CMS field / upload folder", "Display behavior", "Designer export size"],
        cms_rows,
        [1.65, 2.35, 2.25, 2.45, 2.25],
    )
    add_table(
        doc,
        "Static / Placeholder Image Slots",
        ["Image slot", "Where it appears", "Current source", "Display behavior", "Designer export size"],
        static_rows,
        [1.75, 2.1, 2.9, 2.25, 1.95],
    )
    add_table(
        doc,
        "Page-Specific Placement Notes",
        ["Page / section", "Image source", "File reference", "Notes for designer"],
        page_rows,
        [1.9, 2.1, 3.9, 3.05],
        page_break_before=True,
    )
    add_table(
        doc,
        "Quick Designer Checklist",
        ["#", "Asset group", "Required export"],
        checklist_rows,
        [0.45, 3.4, 7.1],
        page_break_before=True,
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Autocracy Machinery - Website Image Resolution Inventory"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(footer, size=8, color="666666")

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_doc())
